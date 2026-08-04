# -*- coding: utf-8 -*-
"""Build a proper Word manual from IQC-操作手冊-v01.md.
Run: python docs/build_manual_docx.py
"""
import re, io, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD = 'docs/IQC-操作手冊-v01.md'
OUT = 'docs/IQC-操作手冊-v01.docx'
FONT = 'Microsoft JhengHei'
GRAY = RGBColor(0x88, 0x88, 0x88)
DARK = RGBColor(0x40, 0x40, 0x40)
ACCENT = RGBColor(0xC0, 0x5A, 0x2B)


def _el(tag):
    return OxmlElement(tag)


def cjk(run):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts if rpr.rFonts is not None else rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), FONT)
    return run


def add_runs(p, text, size=None, color=None, italic=False, bold_all=False):
    """Render text, converting inline **bold** to real bold runs."""
    for seg in re.split(r'(\*\*.+?\*\*)', text):
        if not seg:
            continue
        b = bold_all
        t = seg
        if seg.startswith('**') and seg.endswith('**'):
            t = seg[2:-2]
            b = True
        r = cjk(p.add_run(t))
        r.bold = b
        if size:
            r.font.size = size
        if color:
            r.font.color.rgb = color
        if italic:
            r.italic = True
    return p


def shade_paragraph(p, fill, bar):
    """Callout box: light fill + left accent bar + thin border."""
    pPr = p._p.get_or_add_pPr()
    shd = _el('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill); pPr.append(shd)
    pbdr = _el('w:pBdr')
    left = _el('w:left'); left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18'); left.set(qn('w:space'), '8'); left.set(qn('w:color'), bar); pbdr.append(left)
    for edge in ('top', 'bottom', 'right'):
        e = _el('w:' + edge); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4'); e.set(qn('w:space'), '6'); e.set(qn('w:color'), 'E5E5E5'); pbdr.append(e)
    pPr.append(pbdr)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.12); pf.right_indent = Inches(0.08)
    pf.space_before = Pt(6); pf.space_after = Pt(6)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for typ, txt in (('begin', None), ('instr', 'TOC \\o "1-1" \\h \\z \\u'), ('sep', None)):
        if typ == 'instr':
            it = _el('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = txt; run._r.append(it)
        else:
            fc = _el('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin' if typ == 'begin' else 'separate'); run._r.append(fc)
    t = _el('w:t'); t.text = '（在 Word 中於此處按右鍵 › 更新功能變數，或按 F9 產生目錄）'
    run._r.append(t)
    end = _el('w:fldChar'); end.set(qn('w:fldCharType'), 'end'); run._r.append(end)
    cjk(run); run.font.color.rgb = GRAY; run.italic = True


def add_figure(doc, caption):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.tblPr
    w = _el('w:tblW'); w.set(qn('w:w'), '5000'); w.set(qn('w:type'), 'pct'); tblPr.append(w)
    bd = _el('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = _el('w:' + edge); e.set(qn('w:val'), 'dashed'); e.set(qn('w:sz'), '6'); e.set(qn('w:color'), 'BBBBBB'); bd.append(e)
    tblPr.append(bd)
    cell = tbl.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    va = _el('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)
    shd = _el('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F7F7F7'); tcPr.append(shd)
    tr = tbl.rows[0]
    trPr = tr._tr.get_or_add_trPr()
    h = _el('w:trHeight'); h.set(qn('w:val'), '1900'); h.set(qn('w:hRule'), 'atLeast'); trPr.append(h)
    cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cjk(cp.add_run('（此處插入截圖）')); r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(9)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2); cap.paragraph_format.space_after = Pt(10)
    r = cjk(cap.add_run(caption)); r.font.size = Pt(9); r.font.color.rgb = GRAY


def main():
    src = io.open(MD, 'r', encoding='utf-8').read().splitlines()
    doc = Document()

    # Base style + spacing
    normal = doc.styles['Normal']
    normal.font.name = FONT; normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = normal.paragraph_format; pf.line_spacing = 1.25; pf.space_after = Pt(4)
    # Heading 1 look
    h1 = doc.styles['Heading 1']
    h1.font.name = FONT; h1.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    h1.font.size = Pt(16); h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    in_cover = True
    skip_toc = False
    for s in src:
        s = s.rstrip()
        if in_cover:
            if s.startswith('# '):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(60 if not p._p.getprevious() else 6)
                r = cjk(p.add_run(s[2:].strip())); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            elif s.startswith('**') and s.endswith('**'):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                add_runs(p, s, size=Pt(12), color=DARK)
            elif s.startswith('_') and s.endswith('_'):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cjk(p.add_run(s.strip('_'))); r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
            elif s.startswith('---'):
                doc.add_page_break(); in_cover = False
            continue

        if s.strip() == '':
            continue
        if s.startswith('## ') and ('目' in s and '錄' in s):
            r = cjk(doc.add_heading(level=1).add_run('目錄'))
            add_toc(doc); skip_toc = True
            continue
        if s.startswith('---'):
            if skip_toc:
                doc.add_page_break(); skip_toc = False
            continue
        if skip_toc and s.startswith('- '):
            continue
        if s.startswith('## '):
            cjk(doc.add_heading(level=1).add_run(s[3:].strip()))
        elif s.startswith('**● '):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
            add_runs(p, s)
        elif s.startswith('> '):
            add_figure(doc, s[2:].strip())
        elif s.startswith('**【'):
            p = doc.add_paragraph(); add_runs(p, s, size=Pt(9.5), color=DARK)
            shade_paragraph(p, 'FBF3E7', 'E8A849')
        elif s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, s[2:].strip())
        elif s.startswith('_') and s.endswith('_'):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cjk(p.add_run(s.strip('_'))); r.italic = True; r.font.color.rgb = GRAY
        else:
            add_runs(doc.add_paragraph(), s)

    # Tell Word to recalculate fields (i.e. build the TOC) when the file opens,
    # so the reader doesn't have to press F9 manually.
    upd = _el('w:updateFields'); upd.set(qn('w:val'), 'true')
    doc.settings.element.insert(0, upd)

    try:
        doc.save(OUT)
        print('OK ->', OUT, '| paragraphs:', len(doc.paragraphs))
    except PermissionError:
        alt = OUT.replace('.docx', '-new.docx')
        doc.save(alt)
        print('主檔被開啟中，改存 ->', alt, '（請關閉 Word 後將此檔改名覆蓋，或重跑腳本）')


if __name__ == '__main__':
    main()
